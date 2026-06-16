from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_ROOT = Path(__file__).resolve().parents[1]
TEMPLATE = PROJECT_ROOT / "data" / "processed" / "CJC-Templet_Word2003_converted.docx"
OUTPUT = PROJECT_ROOT / "data" / "processed" / "ST_QueryEngine_Report_final.docx"
FALLBACK_OUTPUT = PROJECT_ROOT / "data" / "processed" / "ST_QueryEngine_Report_final_expanded.docx"


def set_east_asia_font(run, font_name="宋体"):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def clear_body(doc: Document):
    body = doc._body._element
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(9.5)

    style_tokens = [
        ("Title", 18, "黑体", RGBColor(0, 0, 0)),
        ("Heading 1", 10.5, "黑体", RGBColor(0, 0, 0)),
        ("Heading 2", 10, "黑体", RGBColor(0, 0, 0)),
        ("Heading 3", 9.5, "宋体", RGBColor(0, 0, 0)),
    ]
    for name, size, east_font, color in style_tokens:
        style = styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), east_font)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        p_pr = style._element.pPr
        if p_pr is not None and p_pr.numPr is not None:
            p_pr.remove(p_pr.numPr)


def set_section_columns(section, count=1):
    cols = section._sectPr.xpath("./w:cols")
    cols_el = cols[0] if cols else OxmlElement("w:cols")
    if not cols:
        section._sectPr.append(cols_el)
    for child in list(cols_el):
        cols_el.remove(child)
    cols_el.set(qn("w:space"), "720")
    cols_el.set(qn("w:num"), str(count))
    if count == 1 and qn("w:equalWidth") in cols_el.attrib:
        del cols_el.attrib[qn("w:equalWidth")]


def add_paragraph(doc, text="", style=None, align=None, first_line=True):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    fmt = p.paragraph_format
    fmt.line_spacing = 1.08
    fmt.space_after = Pt(3)
    if first_line and style is None:
        fmt.first_line_indent = Pt(18)
    run = p.add_run(text)
    set_east_asia_font(run)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(3)
    for run in p.runs:
        run.bold = True
        set_east_asia_font(run, "黑体" if level <= 2 else "宋体")
    return p


def add_kv_line(doc, key, value):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    r1 = p.add_run(key)
    r1.bold = True
    set_east_asia_font(r1, "黑体")
    r2 = p.add_run(value)
    set_east_asia_font(r2)
    return p


def add_labeled_items(doc, items):
    """在双栏正文中用紧凑段落替代表格，避免分栏分页时表格被切开。"""
    for label, text in items:
        p = doc.add_paragraph()
        fmt = p.paragraph_format
        fmt.first_line_indent = Pt(0)
        fmt.line_spacing = 1.05
        fmt.space_after = Pt(2)
        r1 = p.add_run(f"{label}：")
        r1.bold = True
        set_east_asia_font(r1, "黑体")
        r2 = p.add_run(text)
        set_east_asia_font(r2)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths=None, font_size=7.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    hdr = table.rows[0].cells
    for idx, text in enumerate(headers):
        hdr[idx].text = text
        hdr[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade_cell(hdr[idx], "D9EAF7")
        set_cell_margins(hdr[idx])
        for p in hdr[idx].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(font_size)
                set_east_asia_font(run, "黑体")

    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cells[idx].text = str(text)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[idx])
            for p in cells[idx].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.font.size = Pt(font_size)
                    set_east_asia_font(run)

    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)

    doc.add_paragraph()
    return table


def add_front_matter(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("基于结构分析与混合路由的半结构化表格智能查询系统")
    title_run.bold = True
    title_run.font.size = Pt(18)
    set_east_asia_font(title_run, "黑体")

    add_paragraph(
        doc,
        "作者：__________    学号：__________    课程：数据库原理与运用",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=False,
    )
    add_paragraph(
        doc,
        "单位：__________    项目链接：https://github.com/woem2436/ST-QueryEngine",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=False,
    )

    add_kv_line(
        doc,
        "摘  要：",
        "现实中的 Excel 表格常包含合并单元格、多级表头、横向键值对和章节式说明，难以直接转换为规则关系表并使用 SQL 查询。本文围绕课程项目“表格数据的智能查询系统”，以 SSTQA-zh 数据集为测试基准，设计并实现了一个离线可复现的半结构化表格问答系统。系统首先对原始工作簿进行结构分析与证据化拆分，生成单元格证据、行证据、表级元数据和局部关系视图；随后根据问题语义将查询路由到关系执行、键值查找、稀疏检索、列表抽取、计数、极值和聚合计算等策略。系统参考 ST-Raptor 对半结构化表格“先恢复结构、再执行操作”的思想，但采用轻量规则和本地 SQLite/JSON 存储，避免课堂环境中对外部 LLM 和向量数据库的强依赖。当前系统在 SSTQA-zh 的 764 条测试样本上取得 55.37% 的整体准确率，其中内容匹配类为 61.16%，数值计算类为 43.95%，语义感知类为 47.15%。实验表明，结构保留、混合存储和可解释路由能够显著提高复杂表格问答的稳定性。",
    )
    add_kv_line(
        doc,
        "关键词：",
        "半结构化表格；自然语言查询；混合存储；查询路由；SQLite；SSTQA-zh",
    )

    en_title = add_paragraph(
        doc,
        "A Structure-Aware Hybrid Query Engine for Semi-Structured Table Question Answering",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        first_line=False,
    )
    for run in en_title.runs:
        run.bold = True

    add_kv_line(
        doc,
        "Abstract: ",
        "This project builds an offline and reproducible question answering system for semi-structured Excel tables in SSTQA-zh. The system converts workbooks into cell evidence, row evidence, table metadata and local relational views, then routes natural-language questions to relational execution, key-value lookup, sparse retrieval, list extraction, counting, extreme-value search and aggregation. Inspired by ST-Raptor's structure-first operation pipeline, the implementation uses lightweight rules with SQLite and JSON indexes so that it can run without mandatory external LLM or vector-database services. On 764 test questions, the current system reaches 55.37% overall accuracy.",
    )
    add_kv_line(
        doc,
        "Key words: ",
        "semi-structured tables; natural language query; hybrid storage; query routing; SQLite; SSTQA-zh",
    )


def add_file_details_appendix(doc):
    appendix_section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_section_columns(appendix_section, 1)

    add_heading(doc, "附录A Python 文件实现说明", 1)
    add_paragraph(
        doc,
        "为了便于复现和维护，本节按文件说明项目中主要 Python 文件的职责、输入输出和实现思路。源码总体遵循“入口调度、结构解析、证据索引、混合执行、结果评测”的分层方式：上层文件负责任务流程，下层文件负责存储、检索、分类和答案匹配，测试与脚本文件用于验证和生成报告。",
    )
    add_table(
        doc,
        ["文件", "功能与实现思路"],
        [
            [
                "src/main.py",
                "命令行入口和系统编排文件。它读取 config.yaml，初始化 TableQuerySystem，并根据参数决定执行单条问答、重建索引或全量评测。评测时逐行读取 test.jsonl，根据 table_id 调用 HybridQueryEngine，随后用 AnswerMatcher 判断预测答案是否正确，并把结果写入 evaluation_report.json。",
            ],
            [
                "src/hybrid_engine.py",
                "系统主查询引擎。启动时加载 TableIndexer 生成的 CellRecord 与 RowRecord，并按 table_id 建立内存索引。回答问题时先尝试 RelationalAnswerEngine；若其置信度不足，再按 lookup、list、count、extreme、sum、avg 等路由使用规则执行器和单元格检索兜底。last_trace 用于 explain 模式展示答案来源。",
            ],
            [
                "src/relational_engine.py",
                "当前准确率提升的核心文件。它把部分半结构化表恢复成 TableView，包括表头、记录行、首列和月份宽表列。随后按问题抽取实体、指标、月份、条件和聚合意图，执行筛选、计数、极值、求和、均值、差值和百分比计算。该模块带有置信度门控，避免低可信关系结果覆盖检索结果。",
            ],
            [
                "src/table_indexer.py",
                "索引构建文件。它遍历 data/raw 下的 Excel 文件，填充合并单元格、提取标题和行列上下文，为每个非空单元格生成 CellRecord，为每行生成 RowRecord。索引同时写入 SQLite 和 JSON：SQLite 便于审计和结构化查询，JSON 便于快速加载和离线检索。",
            ],
            [
                "src/text_utils.py",
                "文本规范化与打分工具库。它负责中文问题清洗、数字抽取、Excel 显示值格式化、单位提取、中文 n-gram token 生成、加权重合度计算和去重。该文件使不同模块可以共享同一套文本匹配逻辑，避免各处重复写字符串处理代码。",
            ],
            [
                "src/parser/excel_parser.py",
                "Excel 解析器。它使用 openpyxl 读取公式缓存后的显示值，填充合并区域，裁剪外围全空行列，并在可识别表头时返回带列名的 DataFrame。其设计目标不是立即得到完美关系表，而是尽量保留原始视觉结构，为后续证据索引和局部关系恢复提供基础。",
            ],
            [
                "src/parser/structure_analyzer.py",
                "轻量结构分析器。它可以抽取前若干行的多级表头、将层级表头展平成列名，并根据数据列内容推断 numeric、datetime 或 text 类型。当前主流程主要使用更强的 TableIndexer 与 RelationalAnswerEngine，但该文件保留了可扩展的结构分析接口。",
            ],
            [
                "src/router/query_router.py",
                "透明规则路由器。它按关键词把问题粗分为 SQL_AGG、KV_LOOKUP、VECTOR_SEARCH、CELL_LOOKUP 等类别，并映射到对应 Agent 名称。它的优点是无需 API、结果可解释；缺点是语义覆盖有限，因此主流程中还叠加了 HybridQueryEngine 的细粒度路由。",
            ],
            [
                "src/router/query_classifier.py",
                "可选 LLM 分类器。若 config.yaml 中启用 LLM 且配置 API key，则尝试使用 langchain_deepseek 调用模型输出 JSON 分类；否则自动回退到 RuleBasedRouter。这样既保留语义分类扩展空间，又不影响离线课堂环境运行。",
            ],
            [
                "src/storage/sql_storage.py",
                "SQLite 存储封装。它可以把 DataFrame 写成数据库表，执行 SQL 查询，获取表名并关闭连接。该模块主要服务 SQLAgent 和早期原型，也体现了课程数据库部分的关系型存储接口。",
            ],
            [
                "src/storage/kv_storage.py",
                "JSON 文件形式的 Key-Value 存储。它提供 set、get、delete、contains、get_all 等基础接口，适合保存表级说明、数据来源、单位和运行元数据等不适合放入关系表的轻量信息。",
            ],
            [
                "src/storage/vector_storage.py",
                "离线检索存储。它保留 VectorStorage 的 public API，但默认不依赖 ChromaDB 或 sentence-transformers，而是把文档转换成 sparse-token 计数，按 token 重合分数排序。这样脚本和测试无需下载模型也可以运行。",
            ],
            [
                "src/agents/sql_agent.py",
                "SQL 聚合 Agent。它读取 SQLite 中的首个表和列名，基于问题中的总和、平均、计数、最大、最小等关键词拼接受控 SQL 聚合语句。它是可解释 SQL 查询路径的简单原型，适合规则表，但不直接处理复杂半结构化布局。",
            ],
            [
                "src/agents/kv_agent.py",
                "KV 查询 Agent。它先尝试问题是否直接包含某个 key；若没有，再用 weighted_overlap 在 key 和 value 文本中找最相关项。适合“数据来源是什么”“单位是什么”等元数据问题。",
            ],
            [
                "src/agents/vector_agent.py",
                "检索增强 Agent。它调用 VectorStorage 搜索相关文档片段，若配置了 LLMAgent 则把片段作为上下文交给 LLM；否则直接返回检索到的片段。该模块体现了向量/语义检索扩展路线，但默认可离线运行。",
            ],
            [
                "src/agents/llm_agent.py",
                "可选 LLM 封装器。它读取 DEEPSEEK_API_KEY 等环境变量；未配置 API key 时返回离线兜底说明，而不是让系统报错。这样报告中可以说明 LLM 是可选增强，而不是系统运行前提。",
            ],
            [
                "src/evaluator/answer_matcher.py",
                "答案匹配器。它先做文本规范化和精确匹配，再做数值容差、百分比换算、列表集合匹配和 LCS/ROUGE-L 风格相似度判断。这样可以兼容“50000”和“50000人”、列表顺序差异、长文本包含答案等情况。",
            ],
            [
                "src/evaluator/statistics.py",
                "评测统计器。它保存每条问题的预测、标准答案、是否正确、类别和路由等信息，计算整体准确率和分类准确率，并可导出为 evaluation_report.json。",
            ],
            [
                "scripts/build_report_docx.py",
                "报告生成脚本。它基于课程模板创建最终 DOCX，设置标题、摘要、正文分栏、表格样式和参考文献，并把当前项目指标、实现细节和文件说明写入报告。后续只要修改脚本文本并重新运行即可得到新版报告。",
            ],
            [
                "scripts/vectorize_tables.py",
                "可选的批量检索准备脚本。它遍历 raw Excel 文件，把每行转换为文本片段并写入 VectorStorage。当前主流程不依赖它，但它展示了如何扩展到向量/语义检索路径。",
            ],
            [
                "tests/test_core_behavior.py",
                "核心行为测试。它验证数值答案不能被子串误判、规则路由能识别计数问题、离线 VectorStorage 能按 token 找到相关文档，是防止关键回归的最小测试集。",
            ],
            [
                "tests/test_parser.py",
                "解析调试脚本。它读取示例 Excel，打印解析后的 DataFrame、表格形状、多级表头和展平列名，主要用于观察 ExcelParser 与 StructureAnalyzer 的效果。",
            ],
            [
                "tests/test_router.py",
                "路由调试脚本。它用若干自然语言问题打印 RuleBasedRouter 的分类和目标 Agent，方便检查关键词路由是否符合预期。",
            ],
            [
                "tests/test_storage.py",
                "存储调试脚本。它分别演示 SQLStorage、KVStorage 和 VectorStorage 的基本读写与查询，用于快速确认三类存储接口是否可用。",
            ],
            [
                "tests/test_classifier.py",
                "分类器调试脚本。它初始化 QueryClassifier，对聚合、元数据、列表和复杂分析问题进行分类，用于检查 LLM 未启用时的规则回退表现。",
            ],
            [
                "tests/prepare_kv_data.py",
                "KV 示例数据准备脚本。它向 metadata.json 写入数据来源、表格说明、单位和创建时间等示例元数据，便于测试 KV 查询路径。",
            ],
            [
                "tests/prepare_vector_data.py",
                "检索示例数据准备脚本。它把示例 Excel 的每一行转换为文档片段并写入 VectorStorage，用于演示离线 sparse-token 检索流程。",
            ],
            [
                "各级 __init__.py",
                "包初始化文件。它们主要用于让 src、parser、storage、router、agents、evaluator 等目录被 Python 识别为包，方便模块间通过 import 引用。",
            ],
        ],
        [1.85, 4.65],
        font_size=7.3,
    )


def add_body(doc):
    body_section = doc.add_section(WD_SECTION.CONTINUOUS)
    set_section_columns(body_section, 2)

    add_heading(doc, "1 引言", 1)
    add_paragraph(
        doc,
        "电子表格在现实业务中被广泛用于预算、报销、资产、考核、合同和财务报表等场景，但这类表格并不总是标准二维关系表。标题可能跨行跨列，指标层级可能隐含在合并单元格中，明细行与小计行可能混排，问答清单还可能以章节形式组织。若简单地把第一行视为表头并导入数据库，很多布局语义会被丢失，后续 SQL 查询也难以表达。",
    )
    add_paragraph(
        doc,
        "SIGMOD 2025 论文 ST-Raptor 指出，半结构化表格问答需要同时处理内容定位、表头层级、隐式关系和操作推理。该工作提出 HO-Tree、树操作流水线和前后向验证机制，并发布包含 102 张真实表格、764 个问题的 SSTQA 基准。本文不直接复现大型 LLM 框架，而是在课程项目范围内实现一个轻量、可解释、可评测的系统，重点探索表格拆分、存储选择和查询路由三件事。",
    )

    add_heading(doc, "2 任务定义与数据集", 1)
    add_paragraph(
        doc,
        "项目输入包括 data/raw 目录下的 102 个 xlsx 文件以及 test.jsonl。每条测试样本包含 table_id、自然语言问题、标准答案、问题类型和难度。系统在评测时根据 table_id 读取对应工作簿，返回答案后与标准答案进行文本和数值匹配，并输出整体准确率和分类型准确率。",
    )
    add_labeled_items(
        doc,
        [
            ("表格数量", "102 张真实 Excel 表格"),
            ("问题数量", "764 条自然语言查询"),
            ("问题类型", "内容匹配、数值计算、语义感知"),
            ("评测输出", "evaluation_report.json"),
        ],
    )

    add_heading(doc, "3 系统总体架构", 1)
    add_paragraph(
        doc,
        "系统采用“结构解析-混合存储-查询路由-答案执行-评测统计”的流水线。解析阶段尽量不把表格过早压扁为单一 DataFrame，而是保留坐标、上方表头、左侧标签、同行文本和表级标题。存储阶段同时写入 SQLite 与 JSON 索引。查询阶段先尝试局部关系执行器，置信度不足时再回退到原有混合检索流程。",
    )
    add_paragraph(
        doc,
        "src/main.py 是命令行入口，负责索引重建、单条查询和全量评测；src/table_indexer.py 构建 CellRecord、RowRecord 与 table_index.json；src/relational_engine.py 是本次优化后的核心，负责把部分表格恢复成局部关系视图并执行筛选、计数、极值、求和、均值和差值；src/hybrid_engine.py 负责整合关系执行器、单元格检索和回退策略；src/parser、src/storage、src/router、src/evaluator 分别处理 Excel 解析、存储封装、问题分类和答案评测。",
    )

    add_heading(doc, "4 表格拆分策略", 1)
    add_paragraph(
        doc,
        "拆分策略遵循“能关系化的关系化，不能关系化的证据化”的原则。首先用 openpyxl 读取公式缓存后的显示值，填充所有合并单元格，并裁剪全空边界。随后识别表头候选行、标题行和章节行，为每个非空单元格生成包含 value、row_text、above_text、left_text、right_text、坐标和单位的证据记录。这样即使表格无法直接转成 SQL 表，系统仍能基于上下文定位答案。",
    )
    add_paragraph(
        doc,
        "对于规则明细表，系统构造局部关系视图：将多级表头压缩成列名，过滤小计、合计和说明行，保留行号与原始证据；对于横向月份表，系统识别“1月、2月、3月”等宽表列并支持跨月份运算；对于章节式清单，系统记录当前章节标题，使“某章节一共有多少条目”“涉及几个责任部门”等问题可以在章节范围内统计；对于无显式表头的文本表，系统根据行内模式生成虚拟列，避免把整行文本误当成一个值。",
    )

    add_heading(doc, "5 存储方案选择依据", 1)
    add_paragraph(
        doc,
        "根据题目要求，系统没有把所有表格强行放入一种数据库，而是按数据形态选择存储。SQLite 用于存储单元格、行证据和局部关系结果，优点是可查询、可审计、便于按 table_id 与坐标回溯；JSON 用于保存轻量 n-gram/BM25 风格索引，优点是无需下载 embedding 模型即可离线召回；KV 元数据用于保存文件路径、表格规模、标题和运行配置；向量数据库与 LLM Agent 保留接口，但默认不参与运行，以保证课堂环境可复现。",
    )
    add_labeled_items(
        doc,
        [
            ("单元格与行证据", "写入 SQLite，用于结构化查询、坐标回溯和结果审计。"),
            ("稀疏检索索引", "写入 JSON，使用中文 n-gram/BM25 风格重合度离线召回。"),
            ("表级元数据", "使用 KV/JSON 保存标题、路径、规模、章节和运行配置。"),
            ("语义扩展接口", "保留 LLM/向量检索入口，但默认离线运行，保证课程环境可复现。"),
        ],
    )

    add_heading(doc, "6 查询路由逻辑", 1)
    add_paragraph(
        doc,
        "查询路由分为两层。第一层由 RuleBasedRouter 根据问题中的“多少、最高、最低、总和、平均、哪些、是否、相比”等词汇识别意图。第二层由 HybridQueryEngine 决定具体执行器：若表格适合关系化且置信度足够，则优先调用 RelationalAnswerEngine；否则使用 lookup、list、count、extreme、sum、avg 等轻量执行器和单元格检索兜底。",
    )
    add_paragraph(
        doc,
        "RelationalAnswerEngine 的关键是先抽取条件，再选择目标列。例如“标准分为 2 的考核项目有多少条”会被解析为条件列“标准分=2”和目标对象“考核项目”；“6 月份 A 产品的利润相比 5 月份减少了多少”会被解析为实体 A 产品、指标利润、月份 6 月与 5 月，并执行差值计算。执行后系统返回答案和置信度，低置信度结果不会强行覆盖原检索结果。",
    )

    add_heading(doc, "7 功能完善与 Bug 修复", 1)
    add_paragraph(
        doc,
        "早期版本主要依赖单元格检索，因此在简单查值题上可以工作，但对计数、列表、跨月份运算和多条件筛选较弱。本次优化加入了关系执行器、表格适配性判断、章节范围统计、宽表月份解析、百分比/差值计算、无表头文本表恢复和更保守的列表答案匹配，使系统从“能找相似文本”逐步变成“能执行局部表格操作”。",
    )
    add_labeled_items(
        doc,
        [
            ("指标同步", "将旧报告中的 30.37% 更新为当前全量评测 55.37%。"),
            ("关系表缺失", "新增局部关系视图，让可关系化表格进入筛选和聚合流程。"),
            ("章节计数失败", "记录章节与条目边界，支持章节范围内的 count/list 查询。"),
            ("月份宽表错误", "识别横向月份列，执行跨月差值、百分比和聚合计算。"),
            ("列表答案过严", "增加保守集合匹配，降低顺序和分隔符差异造成的误判。"),
        ],
    )

    add_heading(doc, "8 实验设置与结果", 1)
    add_paragraph(
        doc,
        "实验在本地 Windows + Python 环境中完成，默认不启用在线 LLM API。评测命令为 python src/main.py --evaluate --quiet，系统读取完整 test.jsonl，并将结果写入 data/processed/evaluation_report.json。当前版本全量样本准确率为 55.37%，已接近课程项目设定的 60% 左右目标。",
    )
    add_labeled_items(
        doc,
        [
            ("版本演进", "初始版 30.37%，以单元格检索为主；中间版 46.99%，加入关系执行；当前版 55.37%，增强宽表、章节和列表处理。"),
            ("分类结果", "Content Match：484 条，61.16%；Numeric Computation：157 条，43.95%；Semantic-Aware：123 条，47.15%。"),
        ],
    )
    add_paragraph(
        doc,
        "从路由统计看，764 条问题中有 547 条由 relational 路径处理，184 条由 lookup 路径处理，count、list、extreme 和 sum 路径承担较小但关键的补充作用。结果说明，半结构化表格问答不能只依靠文本相似度；一旦能够恢复局部表结构，数值计算和多条件筛选的上限会明显提高。",
    )

    add_heading(doc, "9 创新性分析", 1)
    add_paragraph(
        doc,
        "本项目的创新性不在于调用更大的模型，而在于把课程数据库思想和半结构化表格特点结合起来。第一，系统以证据单元格和局部关系视图共同表示表格，兼顾 SQL 式精确计算与布局语义保留。第二，查询路由不是 SQL、KV、向量或 LLM 的单选题，而是按问题语义组合不同执行器。第三，系统引入置信度门控，只有当关系执行结果足够可信时才覆盖回退检索，从而减少规则误伤。第四，默认离线可复现，适合课程验收，也为后续接入 LLM 验证器留下空间。",
    )
    add_paragraph(
        doc,
        "与 ST-Raptor 相比，本项目没有构建完整 HO-Tree 和 LLM 操作流水线，而是实现了一个轻量近似：用表头、章节、左右上下文和关系视图表达隐式结构，用规则执行器模拟 lookup、filter、aggregate、compare 等基本操作，再用答案匹配器做结果验证。这一设计牺牲了一部分泛化能力，但换来了实现完整度、可解释性和运行稳定性。",
    )

    add_heading(doc, "10 错误分析与改进方向", 1)
    add_paragraph(
        doc,
        "当前错误主要来自三类样本。第一类是复杂财务表，问题需要跨多个小计区、左右分栏或资产/负债双侧区域进行定位，局部关系视图仍可能选错列。第二类是语义归纳题，例如“总结主要变化”“如何影响整体指数”，需要先抽取多条证据再组织自然语言答案。第三类是数据本身或文件索引存在疑似不一致的样本，系统即使检索正确文件也可能找不到问题所需信息。",
    )
    add_paragraph(
        doc,
        "后续可以沿三条路线继续优化：其一，引入 HO-Tree 或类似树结构，更系统地表达合并单元格和多级表头；其二，增加受限 pandas/SQL 代码生成器，让数值计算题先生成可验证表达式再执行；其三，引入轻量 LLM 验证器，只让模型阅读候选证据而不是整张表，负责目标列选择、条件消歧和摘要生成，以控制成本并减少幻觉。",
    )

    add_heading(doc, "11 结论", 1)
    add_paragraph(
        doc,
        "本文完成了一个面向 SSTQA-zh 的表格数据智能查询系统，实现了从 Excel 结构解析、混合存储、自然语言查询、路由执行到准确率评测的完整闭环。系统当前在 764 条测试样本上达到 55.37% 的准确率，说明结构分析、存储选择和查询路由是解决半结构化表格问答的有效路径。虽然距离 ST-Raptor 这类完整 LLM 框架仍有差距，但该系统已经具备课程项目所需的完整性、可解释性和可继续研究的扩展空间。",
    )

    add_file_details_appendix(doc)

    add_heading(doc, "参考文献", 1)
    references = [
        "[1] Tang Z, Niu B, Zhou X, Li B, Zhou W, Wang J, Li G, Zhang X, Wu F. ST-Raptor: LLM-Powered Semi-Structured Table Question Answering. arXiv:2508.18190, 2025.",
        "[2] OpenDataBox/ST-Raptor. SSTQA-zh dataset. https://github.com/OpenDataBox/ST-Raptor/tree/master/data/SSTQA-zh",
        "[3] weAIDB/ST-Raptor. LLM-Powered Semi-Structured Table Question Answering. https://github.com/weAIDB/ST-Raptor",
        "[4] Python Software Foundation. sqlite3: DB-API 2.0 interface for SQLite databases.",
        "[5] openpyxl documentation. https://openpyxl.readthedocs.io/",
    ]
    for ref in references:
        add_paragraph(doc, ref, first_line=False)


def build_report():
    doc = Document(str(TEMPLATE))
    clear_body(doc)
    configure_styles(doc)

    front_section = doc.sections[0]
    front_section.top_margin = Inches(0.82)
    front_section.bottom_margin = Inches(0.82)
    front_section.left_margin = Inches(0.82)
    front_section.right_margin = Inches(0.82)
    set_section_columns(front_section, 1)

    add_front_matter(doc)
    add_body(doc)

    try:
        doc.save(str(OUTPUT))
        return OUTPUT
    except PermissionError:
        doc.save(str(FALLBACK_OUTPUT))
        return FALLBACK_OUTPUT


if __name__ == "__main__":
    print(build_report())
